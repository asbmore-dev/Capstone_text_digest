"""
output_writer.py
Writes the digest to the requested output format.
Supported formats: txt, pdf, docx, rtf, html
Each output file starts with a two-line header:
  Line 1: Original title
  Line 2: Adaptation style
"""

from __future__ import annotations

import os
from datetime import datetime
from pathlib import Path

OUTPUT_DIR = Path(__file__).parent.parent / "output_files"
OUTPUT_DIR.mkdir(exist_ok=True)


def _timestamp() -> str:
    return datetime.now().strftime("%Y%m%d_%H%M%S")


# ── Public dispatcher ─────────────────────────────────────────────────────────

def write_output(header1: str, header2: str, body: str, fmt: str) -> str:
    """
    Write the digest to a file and return the output path.

    Args:
        header1: Original title.
        header2: Style descriptor, e.g. "Style: Modern".
        body:    Digest body text.
        fmt:     One of 'txt', 'pdf', 'docx', 'rtf', 'html'.

    Returns:
        Absolute file path string.
    """
    writers = {
        "txt":  write_txt,
        "pdf":  write_pdf,
        "docx": write_docx,
        "rtf":  write_rtf,
        "html": write_html,
    }
    if fmt not in writers:
        raise ValueError(f"Unsupported format: {fmt}")

    return writers[fmt](header1, header2, body)


# ── Format writers ────────────────────────────────────────────────────────────

def write_txt(header1: str, header2: str, body: str) -> str:
    path = OUTPUT_DIR / f"digest_{_timestamp()}.txt"
    content = f"{header1}\n{header2}\n\n{body}"
    path.write_text(content, encoding="utf-8")
    return str(path)


def write_pdf(header1: str, header2: str, body: str) -> str:
    try:
        from fpdf import FPDF
    except ImportError:
        raise ImportError("Install fpdf2: uv add fpdf2")

    import unicodedata

    def _to_latin1(text: str) -> str:
        """Sanitize Unicode text to Latin-1 for Helvetica compatibility."""
        replacements = {
            "\u2013": "-",    # en dash
            "\u2014": "--",   # em dash
            "\u2018": "'",    # left single quote
            "\u2019": "'",    # right single quote
            "\u201c": '"',    # left double quote
            "\u201d": '"',    # right double quote
            "\u2026": "...",  # ellipsis
            "\u00a0": " ",    # non-breaking space
            "\u2022": "*",    # bullet
            "\u2122": "(TM)",
            "\u00ae": "(R)",
            "\u00a9": "(C)",
        }
        for uni, ascii_eq in replacements.items():
            text = text.replace(uni, ascii_eq)
        text = unicodedata.normalize("NFKD", text)
        return text.encode("latin-1", errors="ignore").decode("latin-1")

    # A4 page: 210mm wide, left/right margins of 20mm each → 170mm usable
    pdf = FPDF(orientation="P", unit="mm", format="A4")
    pdf.set_margins(left=20, top=20, right=20)
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.add_page()

    # Usable width = page width minus left and right margins
    usable_w = pdf.w - pdf.l_margin - pdf.r_margin  # ~170mm

    # Title header
    pdf.set_font("Helvetica", "B", 16)
    pdf.multi_cell(usable_w, 10, _to_latin1(header1))
    pdf.ln(2)

    # Style line
    pdf.set_font("Helvetica", "I", 11)
    pdf.multi_cell(usable_w, 8, _to_latin1(header2))
    pdf.ln(6)

    # Body
    pdf.set_font("Helvetica", size=11)
    pdf.multi_cell(usable_w, 7, _to_latin1(body))

    path = OUTPUT_DIR / f"digest_{_timestamp()}.pdf"
    pdf.output(str(path))
    return str(path)


def write_docx(header1: str, header2: str, body: str) -> str:
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise ImportError("Install python-docx: uv add python-docx")

    doc = Document()

    # Title
    title_para = doc.add_heading(header1, level=1)
    title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Style subtitle
    style_para = doc.add_paragraph(header2)
    run = style_para.runs[0]
    run.italic = True
    run.font.size = Pt(11)

    doc.add_paragraph()  # blank spacer

    # Body
    doc.add_paragraph(body)

    path = OUTPUT_DIR / f"digest_{_timestamp()}.docx"
    doc.save(str(path))
    return str(path)


def write_rtf(header1: str, header2: str, body: str) -> str:
    """Write a minimal but valid RTF 1.5 file."""
    def _esc(text: str) -> str:
        """Escape non-ASCII chars for RTF."""
        out = []
        for ch in text:
            if ord(ch) < 128:
                if ch in ("\\", "{", "}"):
                    out.append("\\" + ch)
                else:
                    out.append(ch)
            else:
                out.append(f"\\u{ord(ch)}?")
        return "".join(out)

    rtf = (
        r"{\rtf1\ansi\deff0"
        r"{\fonttbl{\f0 Times New Roman;}{\f1 Arial;}}"
        r"{\colortbl;\red0\green0\blue0;}"
        "\n"
        r"{\pard\f1\fs32\b " + _esc(header1) + r"\b0\par}"
        "\n"
        r"{\pard\f1\fs22\i " + _esc(header2) + r"\i0\par}"
        "\n"
        r"{\pard\f0\fs24 " + _esc(body).replace("\n", r"\par ") + r"\par}"
        "\n}"
    )

    path = OUTPUT_DIR / f"digest_{_timestamp()}.rtf"
    path.write_text(rtf, encoding="utf-8")
    return str(path)


def write_html(header1: str, header2: str, body: str) -> str:
    # Convert newlines to <br> for HTML
    html_body = body.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    html_body = html_body.replace("\n", "<br>\n")

    html = f"""<!DOCTYPE html>
<html lang="en">
<head>
  <meta charset="UTF-8">
  <meta name="viewport" content="width=device-width, initial-scale=1.0">
  <title>{header1}</title>
  <style>
    body {{ font-family: Georgia, serif; max-width: 780px; margin: 40px auto;
            padding: 0 20px; color: #222; line-height: 1.7; }}
    h1   {{ font-size: 1.8rem; margin-bottom: 4px; }}
    .style-tag {{ font-style: italic; color: #666; margin-bottom: 24px;
                  font-size: 0.95rem; }}
    p    {{ text-indent: 1.5em; }}
  </style>
</head>
<body>
  <h1>{header1}</h1>
  <div class="style-tag">{header2}</div>
  <p>{html_body}</p>
</body>
</html>"""

    path = OUTPUT_DIR / f"digest_{_timestamp()}.html"
    path.write_text(html, encoding="utf-8")
    return str(path)
